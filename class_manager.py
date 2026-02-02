import os
import json
import re
from pypdf import PdfReader
from docx import Document


class ClassManager:
    def __init__(self):
        self.doc_folder = "class_documents"
        self.db_file = "student_classes.json"

        if not os.path.exists(self.doc_folder):
            os.makedirs(self.doc_folder)
            os.makedirs(os.path.join(self.doc_folder, "class_1"), exist_ok=True)

        self.user_classes = {}
        if os.path.exists(self.db_file):
            try:
                with open(self.db_file, 'r') as f:
                    self.user_classes = json.load(f)
            except:
                self.user_classes = {}

    def get_user_class(self, username):
        return self.user_classes.get(username)

    def set_user_class(self, username, class_id):
        self.user_classes[username] = str(class_id)
        with open(self.db_file, 'w') as f:
            json.dump(self.user_classes, f)
        print(f"✅ Assigned {username} to Class {class_id}")

    def get_class_context(self, class_id):
        root_path = os.path.join(self.doc_folder, f"class_{class_id}")
        if not os.path.exists(root_path):
            return ""

        full_text = []
        print(f"📂 Reading documents for Class {class_id} (Scanning all subjects)...")

        for dirpath, dirnames, filenames in os.walk(root_path):
            for filename in filenames:
                if filename.startswith("~$"): continue

                file_path = os.path.join(dirpath, filename)
                rel_path = os.path.relpath(dirpath, root_path)
                subject_label = "" if rel_path == "." else f"[{rel_path}] "

                try:
                    text = ""
                    if filename.lower().endswith('.pdf'):
                        reader = PdfReader(file_path)
                        for page in reader.pages:
                            t = page.extract_text()
                            if t: text += t + "\n"

                    elif filename.lower().endswith('.docx'):
                        doc = Document(file_path)
                        # 1. Read Paragraphs
                        for para in doc.paragraphs:
                            text += para.text + "\n"
                        # 2. 🟢 NEW: Read Tables (Fix for converted PDFs)
                        for table in doc.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    text += cell.text + " "
                                text += "\n"

                    elif filename.lower().endswith('.txt'):
                        with open(file_path, 'r', encoding='utf-8') as f:
                            text = f.read()

                    if text:
                        text = text.strip()
                        if len(text) > 5:
                            header = f"--- Document: {subject_label}{filename} ---"
                            full_text.append(f"{header}\n{text}\n")

                except Exception as e:
                    print(f"⚠️ Error reading {filename}: {e}")

        return "\n".join(full_text)

    def extract_class_number(self, text):
        match = re.search(r'\b(\d+)\b', text)
        if match: return match.group(1)
        text = text.lower()
        if "one" in text or "yek" in text or "یک" in text or "اول" in text: return "1"
        if "two" in text or "do" in text or "دو" in text or "دوم" in text: return "2"
        if "three" in text or "se" in text or "سه" in text or "سوم" in text: return "3"
        if "four" in text or "chahar" in text or "چهار" in text or "چهارم" in text: return "4"
        if "five" in text or "panj" in text or "پنج" in text or "پنجم" in text: return "5"
        return None



